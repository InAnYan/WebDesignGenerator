from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import datetime
from datetime import datetime
import sys
import os
from html2image import Html2Image
from docx.shared import Inches

hti = Html2Image(size=(1366, 768))
pics_count = 0


def main(argv):
    if len(argv) != 6:
        print('Usage: python3 zvit_maker.py group name n path_in path_out')
        exit(1)
        
    conf = {
        "group": argv[1],
        "name": argv[2],
        "n": argv[3],
        "path_in": argv[4],
        "path_out": argv[5]
    }

    hti.temp_path = conf['path_in']

    files = find_files(conf['path_in'])

    make_zvit(conf, files, conf['path_out'])


def find_files(path_in):
    files = []

    for file_name in sorted(os.listdir(path_in)):
        if file_name.endswith('.html'):
            indecies = file_name[:-5].split('_')
            with open(path_in + '/' + file_name) as fin:
                file_src = fin.read()
                files.append((indecies, file_src))

    return files
    

def make_zvit(conf, files, out_path):
    assert conf != {}
    assert files != []
    assert out_path != ''

    doc = Document()

    first_page(doc, conf)
    second_page(doc, conf, files)

    doc.save(out_path)


def second_page(doc, conf, files):
    add_p(doc, 'Практичне заняття №' + str(conf["n"]), centered=True, bold=True, size=16)

    for f in files:
        add_task(doc, f)


def add_task(doc, f):
    indecies = f[0]
    source = f[1]

    add_p(doc, 'Завдання №' + '.'.join(indecies) + ':', bold=True)

    add_p(doc, 'HTML код', bold=True, centered=True, size=16)
    add_p(doc, source)

    add_hspace(doc)

    global pics_count
    
    pic_path = 'tmp_pic_' + str(pics_count) + '.png'
    hti.screenshot(html_str=source, save_as=pic_path)
    pics_count += 1

    doc.add_picture(pic_path, width=Inches(6))
    
    add_p(doc, 'Скриншот веб-сторінки відповідно до завдання', centered=True, size=16)

    add_hspace(doc)


def first_page(doc, conf):
    add_p(doc, "З В І Т", bold=True, centered=True, size=20)

    add_hspace(doc, 4)

    add_p(doc, 'з практичного заняття №' + str(conf["n"]) + ' дисципліни', centered=True, size=16)
    add_p(doc, 'Програмні та інструментальні засоби web- дизайну', centered=True, size=16, bold=True)
    add_p(doc, 'студента гр. ' + conf["group"], centered=True, size=16)
    add_p(doc, conf["name"], centered=True, size=16)
    
    add_hspace(doc, 10)

    add_p(doc, str(datetime.now().year), centered=True, size=16)

    doc.add_page_break()


def add_p(doc, string, **kwargs):
    p = doc.add_paragraph()

    if kwargs.get('centered', False):
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run(string)
    run.bold = kwargs.get('bold', False)
    run.font.size = Pt(kwargs.get('size', 14))

    return run
    

def add_hspace(doc, n=1):
    for i in range(n):
        doc.add_paragraph()

if __name__ == '__main__':
    main(sys.argv)

